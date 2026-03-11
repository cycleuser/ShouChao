"""
Export module for ShouChao.

Supports exporting briefings and articles to multiple formats:
- PDF: Professional documents with formatting
- EPUB: E-reader compatible format
- HTML: Web-ready documents
- DOCX: Microsoft Word documents
- Audio: TTS-generated audio files
"""

import json
import logging
import os
import tempfile
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any, Optional

logger = logging.getLogger(__name__)


@dataclass
class ExportResult:
    """Result of an export operation."""
    success: bool
    output_path: Optional[str] = None
    format: str = ""
    file_size: int = 0
    error: Optional[str] = None
    metadata: dict = field(default_factory=dict)

    def to_dict(self) -> dict:
        return {
            "success": self.success,
            "output_path": self.output_path,
            "format": self.format,
            "file_size": self.file_size,
            "error": self.error,
            "metadata": self.metadata,
        }


class MarkdownExporter:
    """Export content to Markdown format."""

    def export(
        self,
        content: str,
        title: str,
        output_path: str,
        metadata: Optional[dict] = None,
    ) -> ExportResult:
        try:
            path = Path(output_path)
            if not path.suffix:
                path = path.with_suffix(".md")

            front_matter = self._build_front_matter(title, metadata)
            full_content = f"{front_matter}\n{content}"

            path.write_text(full_content, encoding="utf-8")

            return ExportResult(
                success=True,
                output_path=str(path),
                format="markdown",
                file_size=path.stat().st_size,
                metadata={"title": title},
            )
        except Exception as e:
            logger.error(f"Markdown export error: {e}")
            return ExportResult(success=False, format="markdown", error=str(e))

    def _build_front_matter(self, title: str, metadata: Optional[dict]) -> str:
        meta = metadata or {}
        meta["title"] = title
        meta.setdefault("date", datetime.now().strftime("%Y-%m-%d"))

        lines = ["---"]
        for k, v in meta.items():
            if isinstance(v, list):
                lines.append(f"{k}:")
                for item in v:
                    lines.append(f"  - {item}")
            elif isinstance(v, dict):
                lines.append(f"{k}:")
                for sk, sv in v.items():
                    lines.append(f"  {sk}: {sv}")
            else:
                lines.append(f"{k}: {v}")
        lines.append("---\n")
        return "\n".join(lines)


class PDFExporter:
    """Export content to PDF format."""

    def export(
        self,
        content: str,
        title: str,
        output_path: str,
        metadata: Optional[dict] = None,
        style: str = "professional",
    ) -> ExportResult:
        try:
            from weasyprint import HTML, CSS
        except ImportError:
            return ExportResult(
                success=False,
                format="pdf",
                error="weasyprint not installed. Run: pip install weasyprint",
            )

        try:
            path = Path(output_path)
            if not path.suffix:
                path = path.with_suffix(".pdf")

            html_content = self._markdown_to_html(content, title, style)
            css_content = self._get_css(style)

            html = HTML(string=html_content)
            css = CSS(string=css_content)

            html.write_pdf(path, stylesheets=[css])

            return ExportResult(
                success=True,
                output_path=str(path),
                format="pdf",
                file_size=path.stat().st_size,
                metadata={"title": title, "style": style},
            )
        except Exception as e:
            logger.error(f"PDF export error: {e}")
            return ExportResult(success=False, format="pdf", error=str(e))

    def _markdown_to_html(self, content: str, title: str, style: str) -> str:
        try:
            import markdown
            body = markdown.markdown(content, extensions=["tables", "fenced_code"])
        except ImportError:
            body = f"<pre>{content}</pre>"

        return f"""
<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>{title}</title>
</head>
<body>
    <article>
        <h1 class="title">{title}</h1>
        <div class="content">{body}</div>
    </article>
</body>
</html>
"""

    def _get_css(self, style: str) -> str:
        base_css = """
@page {
    size: A4;
    margin: 2cm;
}
body {
    font-family: "Noto Sans CJK SC", "Source Han Sans SC", "Microsoft YaHei", 
                 "PingFang SC", "Hiragino Sans GB", sans-serif;
    font-size: 11pt;
    line-height: 1.6;
    color: #333;
}
.title {
    font-size: 24pt;
    font-weight: bold;
    color: #1a1a1a;
    margin-bottom: 0.5em;
    padding-bottom: 0.3em;
    border-bottom: 2px solid #333;
}
.content {
    text-align: justify;
}
h1, h2, h3 {
    color: #1a1a1a;
    margin-top: 1.5em;
}
h2 {
    font-size: 16pt;
    border-bottom: 1px solid #ddd;
    padding-bottom: 0.2em;
}
h3 {
    font-size: 13pt;
}
a {
    color: #0066cc;
    text-decoration: none;
}
code {
    background: #f4f4f4;
    padding: 0.1em 0.3em;
    border-radius: 3px;
    font-size: 10pt;
}
pre {
    background: #f4f4f4;
    padding: 1em;
    border-radius: 5px;
    overflow-x: auto;
}
table {
    border-collapse: collapse;
    width: 100%;
    margin: 1em 0;
}
th, td {
    border: 1px solid #ddd;
    padding: 0.5em;
    text-align: left;
}
th {
    background: #f4f4f4;
}
blockquote {
    border-left: 3px solid #ddd;
    margin: 1em 0;
    padding-left: 1em;
    color: #666;
}
"""
        return base_css


class EPUBExporter:
    """Export content to EPUB format."""

    def export(
        self,
        content: str,
        title: str,
        output_path: str,
        metadata: Optional[dict] = None,
    ) -> ExportResult:
        try:
            import ebooklib
            from ebooklib import epub
        except ImportError:
            return ExportResult(
                success=False,
                format="epub",
                error="ebooklib not installed. Run: pip install ebooklib",
            )

        try:
            path = Path(output_path)
            if not path.suffix:
                path = path.with_suffix(".epub")

            meta = metadata or {}

            book = epub.EpubBook()
            book.set_identifier(meta.get("id", f"shouchao-{datetime.now().strftime('%Y%m%d%H%M%S')}"))
            book.set_title(title)
            book.set_language(meta.get("language", "zh"))

            if "author" in meta:
                book.add_author(meta["author"])
            else:
                book.add_author("ShouChao")

            chapter = epub.EpubHtml(title=title, file_name="content.xhtml")
            chapter.content = self._markdown_to_html(content, title)

            book.add_item(chapter)
            book.toc = (chapter,)
            book.add_item(epub.EpubNcx())
            book.add_item(epub.EpubNav())

            style = "body { font-family: sans-serif; line-height: 1.6; }"
            nav_css = epub.EpubItem(
                uid="style_nav",
                file_name="style/nav.css",
                media_type="text/css",
                content=style,
            )
            book.add_item(nav_css)

            book.spine = ["nav", chapter]

            epub.write_epub(str(path), book)

            return ExportResult(
                success=True,
                output_path=str(path),
                format="epub",
                file_size=path.stat().st_size,
                metadata={"title": title},
            )
        except Exception as e:
            logger.error(f"EPUB export error: {e}")
            return ExportResult(success=False, format="epub", error=str(e))

    def _markdown_to_html(self, content: str, title: str) -> str:
        try:
            import markdown
            body = markdown.markdown(content, extensions=["tables"])
        except ImportError:
            body = f"<pre>{content}</pre>"

        return f"<h1>{title}</h1>{body}"


class DOCXExporter:
    """Export content to Microsoft Word format."""

    def export(
        self,
        content: str,
        title: str,
        output_path: str,
        metadata: Optional[dict] = None,
    ) -> ExportResult:
        try:
            from docx import Document
        except ImportError:
            return ExportResult(
                success=False,
                format="docx",
                error="python-docx not installed. Run: pip install python-docx",
            )

        try:
            path = Path(output_path)
            if not path.suffix:
                path = path.with_suffix(".docx")

            doc = Document()

            title_para = doc.add_heading(title, 0)
            title_para.alignment = 1

            for line in content.split("\n"):
                line = line.strip()
                if not line:
                    doc.add_paragraph("")
                elif line.startswith("## "):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith("### "):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith("#### "):
                    doc.add_heading(line[5:], level=4)
                elif line.startswith("- ") or line.startswith("* "):
                    doc.add_paragraph(line[2:], style="List Bullet")
                elif line.startswith("1. ") or line.startswith("1) "):
                    doc.add_paragraph(line[3:], style="List Number")
                else:
                    doc.add_paragraph(line)

            doc.save(str(path))

            return ExportResult(
                success=True,
                output_path=str(path),
                format="docx",
                file_size=path.stat().st_size,
                metadata={"title": title},
            )
        except Exception as e:
            logger.error(f"DOCX export error: {e}")
            return ExportResult(success=False, format="docx", error=str(e))


class HTMLExporter:
    """Export content to HTML format."""

    def export(
        self,
        content: str,
        title: str,
        output_path: str,
        metadata: Optional[dict] = None,
    ) -> ExportResult:
        try:
            import markdown
            body = markdown.markdown(content, extensions=["tables", "fenced_code"])
        except ImportError:
            body = f"<pre>{content}</pre>"

        try:
            path = Path(output_path)
            if not path.suffix:
                path = path.with_suffix(".html")

            meta = metadata or {}
            meta.setdefault("date", datetime.now().strftime("%Y-%m-%d"))

            html = f"""<!DOCTYPE html>
<html lang="{meta.get('language', 'zh')}">
<head>
    <meta charset="utf-8">
    <meta name="viewport" content="width=device-width, initial-scale=1">
    <title>{title}</title>
    <style>
        body {{
            font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, 
                         "Helvetica Neue", Arial, "Noto Sans", sans-serif;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
            line-height: 1.6;
            color: #333;
        }}
        h1 {{ border-bottom: 2px solid #333; padding-bottom: 0.3em; }}
        h2 {{ border-bottom: 1px solid #ddd; padding-bottom: 0.2em; }}
        code {{ background: #f4f4f4; padding: 0.2em 0.4em; border-radius: 3px; }}
        pre {{ background: #f4f4f4; padding: 1em; border-radius: 5px; overflow-x: auto; }}
        blockquote {{ border-left: 3px solid #ddd; margin: 1em 0; padding-left: 1em; color: #666; }}
        table {{ border-collapse: collapse; width: 100%; }}
        th, td {{ border: 1px solid #ddd; padding: 0.5em; }}
        th {{ background: #f4f4f4; }}
    </style>
</head>
<body>
    <article>
        <header>
            <h1>{title}</h1>
            <p><small>Generated: {meta['date']}</small></p>
        </header>
        <main>{body}</main>
    </article>
</body>
</html>"""

            path.write_text(html, encoding="utf-8")

            return ExportResult(
                success=True,
                output_path=str(path),
                format="html",
                file_size=path.stat().st_size,
                metadata={"title": title},
            )
        except Exception as e:
            logger.error(f"HTML export error: {e}")
            return ExportResult(success=False, format="html", error=str(e))


class AudioExporter:
    """Export content to audio format using TTS."""

    def __init__(self, engine: str = "edge-tts"):
        self._engine = engine

    def export(
        self,
        content: str,
        title: str,
        output_path: str,
        language: Optional[str] = None,
        voice: Optional[str] = None,
        rate: float = 1.0,
        on_progress: Optional[callable] = None,
    ) -> ExportResult:
        from shouchao.core.tts import TTSEngine

        try:
            path = Path(output_path)
            if not path.suffix:
                path = path.with_suffix(".mp3")

            tts = TTSEngine(preferred_engine=self._engine)

            full_text = f"{title}。{content}"

            result = tts.synthesize_long(
                text=full_text,
                output_path=str(path),
                engine=self._engine,
                voice=voice,
                language=language,
                rate=rate,
                on_progress=on_progress,
            )

            if result.success:
                return ExportResult(
                    success=True,
                    output_path=result.audio_path,
                    format="audio",
                    file_size=Path(result.audio_path).stat().st_size,
                    metadata={"duration": result.duration, "engine": result.engine},
                )
            else:
                return ExportResult(
                    success=False,
                    format="audio",
                    error=result.error,
                )
        except Exception as e:
            logger.error(f"Audio export error: {e}")
            return ExportResult(success=False, format="audio", error=str(e))


class Exporter:
    """
    Unified export interface supporting multiple formats.

    Usage:
        exporter = Exporter()
        result = exporter.export(
            content="# Hello\\nThis is a test.",
            title="Test Document",
            output_path="output.pdf",
            format="pdf",
        )
    """

    def __init__(self):
        self._exporters = {
            "md": MarkdownExporter(),
            "markdown": MarkdownExporter(),
            "pdf": PDFExporter(),
            "epub": EPUBExporter(),
            "docx": DOCXExporter(),
            "html": HTMLExporter(),
            "audio": AudioExporter(),
        }

    @property
    def supported_formats(self) -> list[str]:
        return list(self._exporters.keys())

    def export(
        self,
        content: str,
        title: str,
        output_path: str,
        format: str = "pdf",
        metadata: Optional[dict] = None,
        **kwargs,
    ) -> ExportResult:
        """
        Export content to specified format.

        Args:
            content: Content to export (Markdown format).
            title: Document title.
            output_path: Output file path.
            format: Export format (pdf, epub, html, docx, md, audio).
            metadata: Additional metadata.
            **kwargs: Format-specific options.

        Returns:
            ExportResult with output path and status.
        """
        fmt = format.lower()

        if fmt not in self._exporters:
            return ExportResult(
                success=False,
                format=fmt,
                error=f"Unsupported format: {fmt}. Supported: {self.supported_formats}",
            )

        exporter = self._exporters[fmt]

        return exporter.export(
            content=content,
            title=title,
            output_path=output_path,
            metadata=metadata,
            **kwargs,
        )


def export_content(
    content: str,
    title: str,
    output_path: str,
    format: str = "pdf",
    **kwargs,
) -> ExportResult:
    """
    Convenience function for exporting content.

    Args:
        content: Content to export.
        title: Document title.
        output_path: Output file path.
        format: Export format.
        **kwargs: Additional options.

    Returns:
        ExportResult with status and path.
    """
    exporter = Exporter()
    return exporter.export(
        content=content,
        title=title,
        output_path=output_path,
        format=format,
        **kwargs,
    )